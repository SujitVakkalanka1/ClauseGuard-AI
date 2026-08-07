import io
import fitz  # PyMuPDF
import docx

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Pure extraction function: accepts raw file bytes and filename.
    Returns plain extracted text.
    Supports PDF (.pdf), Word (.docx), and Plain Text (.txt).
    """
    if not file_bytes:
        raise ValueError("File content is empty.")

    lower_filename = filename.lower()

    if lower_filename.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif lower_filename.endswith(".docx"):
        return _extract_docx(file_bytes)
    elif lower_filename.endswith(".txt"):
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1", errors="ignore")
    else:
        raise ValueError(f"Unsupported file format for file: '{filename}'. Supported types: PDF, DOCX, TXT.")


def _extract_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF bytes using PyMuPDF (fitz)."""
    text_chunks = []
    try:
        # Open PDF from memory stream
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            page_text = page.get_text("text")
            if page_text and page_text.strip():
                text_chunks.append(page_text.strip())
        doc.close()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    extracted = "\n\n".join(text_chunks)
    if not extracted.strip():
        raise ValueError("PDF document contains no readable text or is image-only/scanned.")
    return extracted


def _extract_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes using python-docx."""
    text_chunks = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_chunks.append(paragraph.text.strip())
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_chunks.append(row_text)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX file: {str(e)}")

    extracted = "\n\n".join(text_chunks)
    if not extracted.strip():
        raise ValueError("DOCX document contains no text.")
    return extracted
