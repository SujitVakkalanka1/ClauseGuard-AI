import pytest
import fitz
import docx
import io
from app.services.extraction import extract_text_from_file

def test_extract_txt():
    sample_text = "This is a legal agreement between Party A and Party B."
    file_bytes = sample_text.encode("utf-8")
    extracted = extract_text_from_file(file_bytes, "contract.txt")
    assert extracted == sample_text

def test_extract_pdf():
    # Generate in-memory PDF using PyMuPDF
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "SECTION 1. INDEMNIFICATION\nClient agrees to indemnify Provider.")
    pdf_bytes = doc.tobytes()
    doc.close()

    extracted = extract_text_from_file(pdf_bytes, "agreement.pdf")
    assert "INDEMNIFICATION" in extracted
    assert "Client agrees to indemnify" in extracted

def test_extract_docx():
    # Generate in-memory DOCX using python-docx
    doc = docx.Document()
    doc.add_paragraph("SECTION 2. LIMITATION OF LIABILITY")
    doc.add_paragraph("Provider's liability is capped at $100.")
    
    stream = io.BytesIO()
    doc.save(stream)
    docx_bytes = stream.getvalue()

    extracted = extract_text_from_file(docx_bytes, "terms.docx")
    assert "LIMITATION OF LIABILITY" in extracted
    assert "capped at $100" in extracted

def test_unsupported_format():
    with pytest.raises(ValueError, match="Unsupported file format"):
        extract_text_from_file(b"test", "document.xyz")
